"""PDF and Word document export for compliance assessment results."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF, XPos, YPos

from frontend.ui_helpers import score_grade

FRAMEWORK_DISPLAY = {
    "iso42001": "ISO 42001 - AI Management",
    "iso27017": "ISO 27017 - Cloud Security",
    "iso27018": "ISO 27018 - Cloud PII",
}

STATUS_COLORS_PDF = {
    "ADDRESSED": (220, 242, 220),
    "PARTIAL": (255, 243, 205),
    "GAP": (255, 220, 220),
}

STATUS_COLORS_DOCX = {
    "ADDRESSED": "D4EDDA",
    "PARTIAL": "FFF3CD",
    "GAP": "F8D7DA",
}

HEADER_COLOR = (31, 78, 121)
HEADER_HEX = "1F4E79"

_UNICODE_REPLACEMENTS = {
    "\u2014": "-",  # em dash
    "\u2013": "-",  # en dash
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2026": "...",
    "\u00a0": " ",
}


def _pdf_safe_text(text: str) -> str:
    """Convert text to characters supported by built-in Helvetica (latin-1)."""
    for src, dst in _UNICODE_REPLACEMENTS.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class ReportGenerator:
    """Generate downloadable PDF and Word compliance reports."""

    def generate_pdf(
        self, results: dict, frameworks: list[str], input_source: str
    ) -> bytes:
        """Build a multi-page PDF report and return raw bytes."""
        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)

        self._pdf_cover(pdf, frameworks, input_source)
        self._pdf_executive_summary(pdf, results, frameworks)

        for fw in frameworks:
            self._pdf_framework_section(pdf, fw, results[fw])

        return bytes(pdf.output())

    def _pdf_footer(self, pdf: FPDF) -> None:
        pdf.set_y(-12)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(
            0,
            8,
            f"Page {pdf.page_no()} | AI Governance Toolkit",
            align="C",
        )

    def _pdf_cover(self, pdf: FPDF, frameworks: list[str], input_source: str) -> None:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(31, 78, 121)
        pdf.ln(40)
        pdf.cell(0, 12, "AI Governance Compliance Assessment Report", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("Helvetica", "", 12)
        pdf.set_text_color(60, 60, 60)
        pdf.ln(8)
        pdf.cell(
            0,
            8,
            datetime.now().strftime("%B %d, %Y"),
            align="C",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        pdf.ln(12)
        pdf.cell(
            0,
            8,
            _pdf_safe_text(f"Input source: {input_source}"),
            align="C",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

        fw_names = ", ".join(FRAMEWORK_DISPLAY.get(fw, fw) for fw in frameworks)
        pdf.cell(0, 8, f"Frameworks analyzed: {fw_names}", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_y(-30)
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(
            0,
            8,
            "AI Governance Toolkit | github.com/DavidYeti/ai-governance-toolkit",
            align="C",
        )

    def _pdf_executive_summary(
        self, pdf: FPDF, results: dict, frameworks: list[str]
    ) -> None:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(31, 78, 121)
        pdf.cell(0, 10, "Executive Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        col_widths = [45, 28, 28, 28, 28, 28]
        headers = ["Framework", "Total", "Addressed", "Partial", "Gap", "Score"]

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(*HEADER_COLOR)
        pdf.set_text_color(255, 255, 255)
        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 8, header, border=1, fill=True)
        pdf.ln()

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)
        for fw in frameworks:
            s = results[fw]["summary"]
            row = [
                FRAMEWORK_DISPLAY.get(fw, fw),
                str(s["total"]),
                str(s["addressed"]),
                str(s["partial"]),
                str(s["gap"]),
                f"{s['score_pct']}% ({score_grade(s['score_pct'])})",
            ]
            for i, val in enumerate(row):
                pdf.cell(col_widths[i], 8, _pdf_safe_text(val), border=1)
            pdf.ln()

        self._pdf_footer(pdf)

    def _pdf_framework_section(self, pdf: FPDF, fw: str, fw_data: dict) -> None:
        pdf.add_page()
        summary = fw_data["summary"]
        controls = fw_data["controls"]

        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(31, 78, 121)
        title = f"{FRAMEWORK_DISPLAY.get(fw, fw)} - Score: {summary['score_pct']}%"
        pdf.cell(0, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        col_widths = [32, 58, 28, 32, 28]
        headers = ["Control ID", "Control Name", "Status", "Maturity", "Confidence"]

        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(*HEADER_COLOR)
        pdf.set_text_color(255, 255, 255)
        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 7, header, border=1, fill=True)
        pdf.ln()

        pdf.set_font("Helvetica", "", 7)
        for ctrl in controls:
            if pdf.get_y() > 260:
                self._pdf_footer(pdf)
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 8)
                pdf.set_fill_color(*HEADER_COLOR)
                pdf.set_text_color(255, 255, 255)
                for i, header in enumerate(headers):
                    pdf.cell(col_widths[i], 7, header, border=1, fill=True)
                pdf.ln()
                pdf.set_font("Helvetica", "", 7)

            status = ctrl["status"]
            fill = STATUS_COLORS_PDF.get(status, (255, 255, 255))
            pdf.set_fill_color(*fill)
            pdf.set_text_color(30, 30, 30)

            row = [
                ctrl["control_id"][:18],
                ctrl["control_name"][:42],
                status,
                ctrl["maturity"],
                f"{int(ctrl['confidence_score'] * 100)}%",
            ]
            for i, val in enumerate(row):
                pdf.cell(col_widths[i], 7, _pdf_safe_text(val), border=1, fill=True)
            pdf.ln()

        pdf.ln(6)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(31, 78, 121)
        pdf.cell(0, 8, "Evidence Highlights", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        all_snippets: list[tuple[str, str]] = []
        for ctrl in controls:
            for snippet in ctrl.get("evidence_snippets", []):
                all_snippets.append((ctrl["control_id"], snippet))
                if len(all_snippets) >= 5:
                    break
            if len(all_snippets) >= 5:
                break

        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(50, 50, 50)
        if all_snippets:
            for cid, snippet in all_snippets:
                text = _pdf_safe_text(f"[{cid}] {snippet[:200]}")
                pdf.multi_cell(0, 5, text)
                pdf.ln(2)
        else:
            pdf.cell(0, 6, "No evidence snippets found.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        self._pdf_footer(pdf)

    def generate_docx(
        self, results: dict, frameworks: list[str], input_source: str
    ) -> bytes:
        """Build a Word document report and return raw bytes."""
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Helvetica"
        style.font.size = Pt(10)

        title = doc.add_heading("AI Governance Compliance Assessment Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Input source: {input_source}")
        fw_names = ", ".join(FRAMEWORK_DISPLAY.get(fw, fw) for fw in frameworks)
        doc.add_paragraph(f"Frameworks analyzed: {fw_names}")
        doc.add_paragraph()

        doc.add_heading("Executive Summary", level=2)
        summary_table = doc.add_table(rows=1, cols=6)
        summary_table.style = "Table Grid"
        sum_headers = ["Framework", "Total Controls", "Addressed", "Partial", "Gap", "Score"]
        for i, header in enumerate(sum_headers):
            cell = summary_table.rows[0].cells[i]
            cell.text = header
            self._shade_cell(cell, HEADER_HEX)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, fw in enumerate(frameworks):
            s = results[fw]["summary"]
            row = summary_table.add_row().cells
            values = [
                FRAMEWORK_DISPLAY.get(fw, fw),
                str(s["total"]),
                str(s["addressed"]),
                str(s["partial"]),
                str(s["gap"]),
                f"{s['score_pct']}% ({score_grade(s['score_pct'])})",
            ]
            for i, val in enumerate(values):
                row[i].text = val
                if row_idx % 2 == 1:
                    self._shade_cell(row[i], "F0F4F8")

        doc.add_paragraph()

        for fw in frameworks:
            fw_data = results[fw]
            summary = fw_data["summary"]
            controls = fw_data["controls"]

            doc.add_heading(
                f"{FRAMEWORK_DISPLAY.get(fw, fw)} - Score: {summary['score_pct']}%",
                level=2,
            )

            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            ctrl_headers = ["Control ID", "Control Name", "Status", "Maturity", "Confidence"]
            for i, header in enumerate(ctrl_headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                self._shade_cell(cell, HEADER_HEX)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

            for ctrl in controls:
                row = table.add_row().cells
                status = ctrl["status"]
                row[0].text = ctrl["control_id"]
                row[1].text = ctrl["control_name"]
                row[2].text = status
                row[3].text = ctrl["maturity"]
                row[4].text = f"{int(ctrl['confidence_score'] * 100)}%"
                fill = STATUS_COLORS_DOCX.get(status, "FFFFFF")
                for cell in row:
                    self._shade_cell(cell, fill)

            doc.add_heading("Evidence Highlights", level=3)
            all_snippets: list[tuple[str, str]] = []
            for ctrl in controls:
                for snippet in ctrl.get("evidence_snippets", []):
                    all_snippets.append((ctrl["control_id"], snippet))
                    if len(all_snippets) >= 5:
                        break
                if len(all_snippets) >= 5:
                    break

            if all_snippets:
                for cid, snippet in all_snippets:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.add_run(f"[{cid}] ").bold = True
                    p.add_run(snippet[:500])
            else:
                doc.add_paragraph("No evidence snippets found.")

            doc.add_paragraph()

        footer = doc.add_paragraph(
            "Generated by AI Governance Toolkit | github.com/DavidYeti/ai-governance-toolkit"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _shade_cell(self, cell, hex_color: str) -> None:
        """Apply background shading to a Word table cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)
